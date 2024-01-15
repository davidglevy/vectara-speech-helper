from abc import ABC
from vectara_client.domain import ResponseSet
from typing import Any, List, Type
import marko
from marko.block import Document as MarkoDocument, BlankLine, Paragraph, ThematicBreak, List as MarkoList
from marko.inline import Element, InlineElement, RawText, Emphasis, StrongEmphasis, LineBreak
from marko.block import SetextHeading, Heading
from docx import Document as DocXDocument
import logging
import re
from jinja2 import Template
import codecs

PARAM_PATTERN = re.compile(r'(\${[A-Za-z0-9_]+})')

def replace_params(input_text, values: dict[str, Any]):
    pos = 0
    result = ""
    for m in re.finditer(PARAM_PATTERN, input_text):
        start = m.start(1)
        result += input_text[pos:start]

        # Get the param minus wrapping
        param = m.group(1)[2:-1]

        if param in values:
            value = values[param]
            result += value
        else:
            raise Exception(f"Parameter [{param}] not found")

        pos = start + len(m.group(1))
    return result

class BaseRenderer(ABC):
    """
    Base renderer which will create an output base on the supplied markdown templates.


    """
    def __init__(self, template, output):
        self.cover_template = template
        self.output = output
        self.logger = logging.getLogger(__class__.__name__)


    def render(self, value_map:dict[str,Any]):
        # Read the markdown templates.
        with open(self.cover_template, 'r') as f:
            cover_page = f.read()

        # Convert the Jinja Markdown to interpolated with repeats
        template = Template(cover_page, trim_blocks=True)

        rendered_markdown = re.sub(r'[^\x00-\x7F]+',' ',template.render(data=value_map))
        self.logger.info(f"Markdown for presentation is:\n{rendered_markdown}")

        document = marko.parse(rendered_markdown)

        self._render_impl(document)

    def _render_impl(self, doc:MarkoDocument):
        raise Exception("Implement in Subclasses")

class WordElementRenderer(ABC):

    def render(self, to_render:Element, doc:DocXDocument):
        raise Exception("Implement in Subclass")

    def text(self, to_render:Element):
        raise Exception("Implement in Subclass")

    def can_handle(self, element):
        raise Exception("Implement in Subclass")

class RenderElementFactory:



    def __init__(self):
        self.element_renderers = [] #:List[WordElementRenderer]

    def add_renderer(self, to_add: WordElementRenderer):
        self.element_renderers.append(to_add)

    def find_renderer(self, element:Element) -> WordElementRenderer:
        for renderer in self.element_renderers:
            if renderer.can_handle(element):
                return renderer

        raise Exception(f"No element renderer found for element [{element}]")

t = Type['t']

def cast(typ:t, val) -> t:
    return val

class BaseWordElementRenderer(WordElementRenderer):

    def __init__(self, factory:RenderElementFactory):
        self.factory = factory
        self.logger = logging.getLogger(self.__class__.__name__)

class TitleElementRenderer(BaseWordElementRenderer):

    def render(self, to_render:Element, doc:DocXDocument):
        # For each child, iterate with text method
        # Then once we have all the text, create a title
        el = cast(SetextHeading, to_render)

        title = ""
        for child in el.children:
            child_renderer = self.factory.find_renderer(child)
            title += child_renderer.text(child)

        doc.add_heading(title, 0)

    def can_handle(self, element):
        return isinstance(element, SetextHeading)

class HeadingRenderer(BaseWordElementRenderer):

    def render(self, to_render:Element, doc:DocXDocument):
        # For each child, iterate with text method
        # Then once we have all the text, create a title
        el = cast(SetextHeading, to_render)

        title = ""
        for child in el.children:
            child_renderer = self.factory.find_renderer(child)
            title += child_renderer.text(child)

        doc.add_heading(title, level=1)

    def can_handle(self, element):
        return isinstance(element, Heading)


class RawTextRenderer(BaseWordElementRenderer):
    def render(self, to_render:Element, doc:DocXDocument):
        raise Exception("To Be Implemented")

    def text(self, to_render:Element):
        return to_render.children # str

    def can_handle(self, element):
        return isinstance(element, RawText) or isinstance(element, marko.inline.CodeSpan)

class EmphasisRenderer(BaseWordElementRenderer):
    def render(self, to_render:Element, doc:DocXDocument):
        raise Exception("To Be Implemented")

    def text(self, to_render:Element):
        text = ""
        for child in to_render.children:
            child_renderer = self.factory.find_renderer(child)
            text += child_renderer.text(child)
        return text

    def can_handle(self, element):
        return isinstance(element, Emphasis) or isinstance(element, StrongEmphasis)

class BlankLineRenderer(BaseWordElementRenderer):
    def render(self, to_render:Element, doc:DocXDocument):
        doc.add_paragraph('')

    def text(self, to_render:Element):
        return "\n"

    def can_handle(self, element):
        return isinstance(element, BlankLine) or isinstance(element, LineBreak)

class PageBreakRenderer(BaseWordElementRenderer):
    def render(self, to_render:Element, doc:DocXDocument):
        doc.add_page_break()

    def text(self, to_render:Element):
        return "\n"

    def can_handle(self, element):
        return isinstance(element, ThematicBreak)

class ParagraphRenderer(BaseWordElementRenderer):
    def render(self, to_render:Element, doc:DocXDocument):
        # TODO Parse paragraph elements with style
        el = cast(Paragraph, to_render)

        text = ""
        for child in el.children:
            child_renderer = self.factory.find_renderer(child)
            text += child_renderer.text(child)

        doc.add_paragraph(text)

    def text(self, to_render:Element):
        return "\n"

    def can_handle(self, element):
        return isinstance(element, Paragraph) or isinstance(element, marko.block.FencedCode)

class ListRenderer(BaseWordElementRenderer):
    def render(self, to_render:Element, doc:DocXDocument):
        # TODO Parse paragraph elements with style
        el = cast(Paragraph, to_render)

        for child in el.children:
            if isinstance(child, marko.block.ListItem):
                text = ""
                p = doc.add_paragraph('', style='List Number')
                # First child will be a paragraph
                para_child = child.children[0]
                for list_child in para_child.children:
                    child_renderer = self.factory.find_renderer(list_child)
                    p.add_run(child_renderer.text(list_child))
            else:
                raise Exception(f"Child of List is not a ListItem: {child}")

    def text(self, to_render:Element):
        return "\n"

    def can_handle(self, element):
        return isinstance(element, MarkoList)


class WordRenderer(BaseRenderer):

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.logger = logging.getLogger(self.__class__.__name__)
        factory = RenderElementFactory()

        factory.add_renderer(TitleElementRenderer(factory))
        factory.add_renderer(RawTextRenderer(factory))
        factory.add_renderer(BlankLineRenderer(factory))
        factory.add_renderer(ParagraphRenderer(factory))
        factory.add_renderer(EmphasisRenderer(factory))
        factory.add_renderer(HeadingRenderer(factory))
        factory.add_renderer(PageBreakRenderer(factory))
        factory.add_renderer(ListRenderer(factory))

        self.factory = factory



    def _render_impl(self, doc:MarkoDocument):

        output = DocXDocument()

        for child in doc.children:
            renderer = self.factory.find_renderer(child)
            renderer.render(child, output)

        output.save(self.output)

